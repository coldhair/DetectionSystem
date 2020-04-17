import re
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Pt

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
import os
import csv


def parse(Path):
    parser = PDFParser(Path)
    document = PDFDocument(parser)
    re_list = []

    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            layout = device.get_result()
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    results = x.get_text()
                    re_list.append(results)
    print(re_list)
    return re_list


document = Document()

# 以下代码设置页布局为横向，比较麻烦
section = document.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 正则表达式部分
datepat = re.compile(r'(20[1-2]\d)[年\s]{1,3}([0-1][0-9])[月\s]{1,3}([0-3][0-9])')
namedate = re.compile(r'吕艳朋|张高峰|马好|邸欲晓|李铭')
compat = re.compile(r'中国.*分公司')
damaipat = re.compile(r'(0\d{11})\n')  # 发票代码
inv_numpat = re.compile(r'\n(\d{8})\n')
valuepat = re.compile(r'[¥|￥]\s?(\d+\.\d\d)\n')
monthpat = re.compile(r'[帐|账]期:20[1-2]\d.?([0-1][0-9])')

if __name__ == '__main__':
    iterow = []
    headname = ['开票日期', '报销人', '开票单位', '发票代码', '发票号码', '报销类别', '月控制额度', '报销月份', '发票金额', '张数']
    iterow.append(headname)
    table = document.add_table(rows=1, cols=10, style='Table Grid')
    headrow_cells = table.rows[0].cells
    for i in range(10):
        headrow_cells[i].text = headname[i]

    lis = []
    for root, dirs, files in os.walk(r"D:\phone\pdf"):
        for file in files:
            print(file.split('.')[-1])
            if file.split('.')[-1] == 'pdf':
                dir = os.path.join(root, file)
                Path = open(dir, 'rb')
                invoice_info = parse(Path)
                invoice_str = ''.join(invoice_info)
                t = datepat.findall(invoice_str)
                # print(t[0])
                y, m, d = t[0]
                date = y + '.' + m + '.' + d
                lis.append(date)
                nm = namedate.findall(invoice_str)[0]  # 姓名
                lis.append(nm)
                com = compat.findall(invoice_str)[0]
                lis.append(com)
                damai = damaipat.findall(invoice_str)[0]  # 发票代码
                lis.append(damai)
                inv_num = inv_numpat.findall(invoice_str)[0]
                lis.append(inv_num)
                value = valuepat.findall(invoice_str)[0]  # 发票金额
                month = monthpat.findall(invoice_str)[0]
                if nm == '张高峰':
                    item, limit = '科长', '150'
                elif nm == '邸欲晓':
                    item, limit = '其他人员', '50'
                else:
                    item, limit = '机动队员', '70'
                lis.append(item)
                lis.append(limit)
                lis.append(month)
                lis.append(value)
                lis.append('1')
                print(lis)
                row_cells = table.add_row().cells
                for i in range(10):
                    row_cells[i].text = lis[i]
                lis = []
                iterow.append(lis)
    with open('some.csv', 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerows(iterow)

    document.save(r'D:\phone\话费报销信息.docx')
