from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pypinyin import pinyin

with open('hanzi.txt', encoding='utf8') as f:
    str = ''.join(f.readlines())

pinyin_list = pinyin(str, heteronym=True)
x = ''

doc = Document()
doc.styles['Normal'].font.name = u'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.add_paragraph()
p = doc.add_paragraph()
for pinyin in pinyin_list:
    py = ' '.join(pinyin)
    run = p.add_run(py + ' ')
    run.font.size = Pt(21)
    if len(pinyin) > 1:
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

doc.save('result.docx')
