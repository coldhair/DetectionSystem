from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from pypinyin import pinyin

with open('hanzi.txt', encoding='utf8') as f:
    str = ''.join(f.readlines()).strip().strip('\ufeff')
    char_list = str.split()
    print(char_list)

doc = Document()
doc.styles['Normal'].font.name = u'Tw Cen MT'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Tw Cen MT')

doc.sections[0].left_margin = Inches(0.5)  # 设置左侧页边距
doc.sections[0].right_margin = Inches(0.5)
doc.sections[0].top_margin = Inches(0.5)
doc.sections[0].bottom_margin = Inches(0.5)


doc.add_paragraph()
p = doc.add_paragraph()
pp = [pinyin(wor) for wor in char_list]
a = []
for x in pp:
    y = [i[0].ljust(4) for i in x]
    a.append(y)

print(a)

b = [' '.join(i) for i in a]
print(b)
long = 80
runp = []
hang = ''

while len(hang) <= long and b != []:
    add = b.pop(0)
    # run = p.add_run('[' + add + ']  ')
    run = p.add_run('[')
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    run = p.add_run(add)
    run = p.add_run(']  ')
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    hang += '[' + add + ']  '
    runp.append(add)

    if len(hang) >= long:
        p = doc.add_paragraph()
        for piy in runp:
            # run = p.add_run('(' + piy + ')  ')
            run = p.add_run('(')
            run = p.add_run(piy)
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            run = p.add_run(')  ')
        p = doc.add_paragraph()
        hang = ''
        runp = []
if b == [] and len(hang) < long:  # 添加最后一行
    p = doc.add_paragraph()
    for piy in runp:
        run = p.add_run('(')
        run = p.add_run(piy)
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run = p.add_run(')  ')

# 用hex(255)转化为十六进制控制颜色为白色

doc.save('result.docx')

