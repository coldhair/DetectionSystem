from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.shared import Pt
import math
import time
d=time.strftime('%Y年%m月%d日',time.localtime())
print(d)

with open('Meal_ordering.txt', encoding='utf8') as file:
    lines = file.readlines()
    print(lines)
    num = len(lines)
doc = Document()
doc.styles['Normal'].font.name = u'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

for i in range(1):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('应急机动队员领餐签字表')
run.font.size = Pt(21)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.add_run(d+'午餐')
run.font.size = Pt(18)
#
# if num % 13 == 0:
#     columns = num / 13
# else:
#     columns = num // 13 + 1

columns = math.ceil(num / 13)  # 向上取整，跟上面注释掉的是相同效果

table = doc.add_table(rows=14, cols=columns * 3, style='Table Grid')
# table.cell(0, 0).width = Pt(30)  # 设置列宽
# table.cell(0, 1).width = Pt(120)  # 设置列宽
# table.cell(0, 2).width = Pt(30)  # 设置列宽
# table.cell(0, 3).width = Pt(120)  # 设置列宽


for x in range(columns):
    row_cells = table.rows[0].cells
    row_cells[x * 3 + 0].text = '序号'
    row_cells[x * 3 + 1].text = '姓名'
    row_cells[x * 3 + 2].text = '签收'
    for i in range(1, num+1):
        if i % 13 == 0:
            m, y = 13, i // 13 - 1
        else:
            m, y = i % 13, i // 13
        row_cells = table.rows[m].cells
        row_cells[y * 3 + 0].text = str(i)
        row_cells[y * 3 + 1].text = lines[i - 1].strip()
        row_cells[y * 3 + 2].text = '\t'

doc.save('neww.docx')
