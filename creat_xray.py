from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.section import WD_SECTION  # 关于节的格式
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from pagesnum import add_numpages, add_page_number
import pymysql
import cntime
import merger # 同类归并
import tupleliststr

document = Document()
document.styles['Normal'].font.name = u'Times New Roman'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

db = pymysql.connect('localhost', 'root', '6579178', 'report_sys')
cursor = db.cursor()

# Unit_info_sql = "INSERT INTO unit_info(unit_name, unit_address, unit_linkman, unit_phone) VALUES(%s,%s,%s,%s)"


for i in range(4):
    document.add_paragraph()  # 添加空白段落

p = document.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('济南铁路疾病预防控制中心')
run.font.size = Pt(21)

p = document.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('检测与评价报告书')
run.font.size = Pt(42)

# cursor.execute("SELECT LAST_INSERT_ID()")  # 获取最后插入时的ID
# unit_id_ob = cursor.fetchone()
# unit_id = unit_id_ob[0]

# 获取报告书的年份和报告书编号
num_sql = "SELECT DATE_FORMAT(report_date,'%Y'),report_num FROM report_info WHERE report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(num_sql)  # 获取报告书年份
num_ob = cursor.fetchone()
year = num_ob[0]
report_num = num_ob[1]

p = document.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('济铁 疾控检 字（{}）第{}号'.format(year, report_num))
run.font.size = Pt(15.75)

for i in range(6):
    document.add_paragraph()  # 添加空白段落

# 添加段落并定义段落格式
p = document.add_paragraph()
para_format = p.paragraph_format
para_format.first_line_indent = Pt(18)  # 首行缩进
para_format.line_spacing = 1.2  # 1.2倍行距

unit_info_sql = "SELECT unit_name,unit_linkman,unit_phone,unit_address FROM unit_info WHERE unit_id=(SELECT report_unit FROM report_info WHERE report_id=(SELECT MAX(report_id) FROM report_info))"
cursor.execute(unit_info_sql)
unit_info_ob = cursor.fetchone()
unit_name = unit_info_ob[0]
unit_linkman = unit_info_ob[1]
unit_phone = unit_info_ob[2]
unit_address = unit_info_ob[3]

run1 = p.add_run('用 人 单 位：')
run1.font.size = Pt(15.75)
run1 = p.add_run(' {} '.format(unit_name))
run1.underline = WD_UNDERLINE.SINGLE  # 单下划线
run1.font.size = Pt(15.75)

# 添加段落并定义段落格式
p = document.add_paragraph()
para_format = p.paragraph_format
para_format.first_line_indent = Pt(18)  # 首行缩进
para_format.line_spacing = 1.2  # 1.2倍行距

run1 = p.add_run('项 目 名 称：')
run1.font.size = Pt(15.75)
run1 = p.add_run(' X射线行李包检查系统卫生防护检测 ')
run1.underline = WD_UNDERLINE.SINGLE  # 单下划线
run1.font.size = Pt(15.75)

for i in range(4):
    document.add_paragraph()  # 添加空白段落

p2 = document.add_paragraph()
p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

report_date_sql = "SELECT UNIX_TIMESTAMP(report_date),report_num FROM report_info WHERE report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(report_date_sql)
report_date_ob = cursor.fetchone()
report_date_timestamp = report_date_ob[0]
report_date = cntime.chinese_data(report_date_timestamp)
run1 = p2.add_run('报告日期   ')
run1.font.size = Pt(15.75)
run1 = p2.add_run('{}'.format(report_date))
run1.font.size = Pt(15.75)

# document.add_page_break() #添加分节后不需要加分页符也会产生新的一页
current_section = document.add_section()
current_section.start_type = WD_SECTION.NEW_PAGE

header = document.sections[1].header
header.is_linked_to_previous = False
hdp1 = header.paragraphs[0]
hdp1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
hrun1 = hdp1.add_run('济铁 疾控检 字（{}） 第{}号'.format(year, report_num))
hrun1.underline = WD_UNDERLINE.SINGLE
hrun1.font.size = Pt(9)

hdp2 = header.add_paragraph()
hrun2 = hdp2.add_run('检 测 与 评 价 报 告 首 页')
hrun2.font.size = Pt(15.75)
hrun2 = hdp2.add_run(' \t\t\t\t\t\t共')
hrun2.font.size = Pt(9)
add_numpages(hrun2)  # 加上页码
hrun2 = hdp2.add_run('页，第')
hrun2.font.size = Pt(9)
add_page_number(hrun2)
hrun2 = hdp2.add_run('页')
hrun2.font.size = Pt(9)

# 添加段落底部边框
b = hdp2._p
pPr = b.get_or_add_pPr()  # 获取段落属性
bottomborder = OxmlElement('w:bottom')  # 添加段落底部边框
bottomborder.set(qn('w:val'), "single")  # 线型
pPr.append(bottomborder)

# 添加页脚
footer = document.sections[1].footer
footer.is_linked_to_previous = False
footp1 = footer.paragraphs[0]
frun1 = footp1.add_run('检测报告包括：封面、首页、正文（附页）、封底，并盖有计量认证章、检测章和骑缝章。')
frun1.font.size = Pt(9)

# 添加页脚段落顶部边框
b = footp1._p
pPr = b.get_or_add_pPr()  # 获取段落属性
topborder = OxmlElement('w:top')  # 添加段落底部边框
topborder.set(qn('w:val'), "single")  # 线型
pPr.append(topborder)
footer.add_paragraph()

# 以下是第二节的内容，报告首页
# 插入表格
table = document.add_table(rows=4, cols=4, style="Table Grid")  # 一次生成4x4表效率更高
table.autofit = True
table.cell(1, 0).width = Pt(30)  # 设置列宽
table.cell(1, 1).width = Pt(260)  # 设置列宽
table.cell(1, 2).width = Pt(50)  # 设置列宽
table.cell(1, 3).width = Pt(160)  # 设置列宽
for i in range(4):
    table.rows[i].height = Cm(0.8)  # 设置行高

headrow_cells = table.rows[0].cells
headrow_cells[0].text = '检测项目：空气比释动能率'  # 此处应从数据库中读取
row_cells = table.rows[1].cells  # 添加第二行
print(row_cells)
row_cells[0].text = '通信地址'
row_cells[1].text = '{}'.format(unit_address)
row_cells[2].text = '联 系 人'
row_cells[3].text = "{}".format(unit_linkman)

test_day_sql = "SELECT DATE_FORMAT(MIN(test_day),'%Y年%m月%d日'),DATE_FORMAT(MAX(test_day),'%Y年%m月%d日') FROM xray_info WHERE report_id=(SELECT MAX(report_id) FROM report_info);"
cursor.execute(test_day_sql)
test_day_ob = cursor.fetchone()
test_minday, test_maxday = test_day_ob[0], test_day_ob[1]

row_cells = table.rows[2].cells  # 添加第二行
row_cells[0].text = '检测时间'
if test_minday == test_maxday:  # 判断一下是否是同一天检测的
    row_cells[1].text = '{}'.format(test_minday)
else:
    row_cells[1].text = '{}—{}'.format(test_minday, test_maxday)

row_cells[2].text = '电    话'
row_cells[3].text = "{}".format(unit_phone)

instrument_sql = "SELECT instrument_name,instrument_num FROM instrument WHERE instrument_id=(SELECT MAX(instrument_id) FROM instrument);"
cursor.execute(instrument_sql)
instrument_ob = cursor.fetchone()
instrument_name = instrument_ob[0]
instrument_num = instrument_ob[1]

row_cells = table.rows[3].cells  # 添加第四行
row_cells[0].text = '现场仪器'
row_cells[1].text = '{}（{}）'.format(instrument_name, instrument_num)

table.cell(0, 0).merge(table.cell(0, 3))  # 合并表格
table.cell(3, 1).merge(table.cell(3, 3))  # 合并表格

# 添加检测依据部分
p = document.add_paragraph()
# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
run = p.add_run('检测与评价依据:')
run.font.size = Pt(11)
# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
run = p.add_run('GBZ 127－2002  X射线行李包检查系统卫生防护标准')
run.font.size = Pt(11)
# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
run = p.add_run('GB 18871-2002  电离辐射防护与辐射源安全基本标准')
run.font.size = Pt(11)
# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
run = p.add_run('（本栏以下空白）')
run.font.size = Pt(11)
# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

for i in range(13):
    document.add_paragraph()  # 添加空白段落

# 添加段落底部边框
p = document.add_paragraph()
b = p._p
pPr = b.get_or_add_pPr()  # 获取段落属性
bottomborder = OxmlElement('w:bottom')  # 添加段落底部边框
bottomborder.set(qn('w:val'), "single")  # 线型
pPr.append(bottomborder)

# 设置行间距
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run('报告编制：')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.add_run('批准：\t\t\t\t\t\t盖章\t\t')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.line_spacing = 1

p = document.add_paragraph()
run = p.add_run('审核：\t\t\t\t\t\t\t\t\t{}'.format(report_date))
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.line_spacing = 1

# 添加第二个分节符
current_section = document.add_section()
current_section.start_type = WD_SECTION.NEW_PAGE

header = document.sections[2].header
header.is_linked_to_previous = False
hdp1 = header.paragraphs[0]
hdp1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
hrun1 = hdp1.add_run('济铁 疾控检 字（{}） 第{}号'.format(year, report_num))
hrun1.underline = WD_UNDERLINE.SINGLE
hrun1.font.size = Pt(9)

hdp2 = header.add_paragraph()
# hdp2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
hrun2 = hdp2.add_run('检 测 与 评 价 报 告')
hrun2.font.size = Pt(15.75)
hrun2 = hdp2.add_run(' \t\t\t\t\t\t\t共')
hrun2.font.size = Pt(9)
add_numpages(hrun2)  # 加上页码试试效果
hrun2 = hdp2.add_run('页，第')
hrun2.font.size = Pt(9)
add_page_number(hrun2)
hrun2 = hdp2.add_run('页')
hrun2.font.size = Pt(9)

b = hdp2._p
pPr = b.get_or_add_pPr()  # 获取段落属性
bottomborder = OxmlElement('w:bottom')  # 添加段落底部边框
bottomborder.set(qn('w:val'), "single")  # 线型
pPr.append(bottomborder)

# 正文部分
p = document.add_paragraph()
run = p.add_run('一、检测评价范围')
run.font.size = Pt(12)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run(
    '{}各客运车站主要生产任务是旅客的运输和行李货物的运输。'.format(unit_name) +
    'X射线行李包检查系统主要是检查行李包裹中的危险品，其职业有害因素是检查行李包裹过程中产生的X射线，'
    '检测其空气比释动能率，并对其防护效果做出评价。')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('二、工作场所情况（工作场所及岗位情况、工作情况及个体防护情况）')
run.font.size = Pt(12)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('{}行包检测仪分布情况如下表'.format(unit_name))
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

headname = ['站名', '型号', '安装日期', '安装位置', '接触人数', '工作班制']
table = document.add_table(rows=1, cols=6, style='Table Grid')
table.cell(0, 0).width = Pt(30)  # 设置列宽
table.cell(0, 1).width = Pt(120)  # 设置列宽
table.cell(0, 2).width = Pt(30)  # 设置列宽
table.cell(0, 3).width = Pt(120)  # 设置列宽

headrow_cells = table.rows[0].cells
for i in range(6):
    headrow_cells[i].text = headname[i]

xray_info_sql = "SELECT station,product_model, DATE_FORMAT(service_time,'%Y.%m'),location,exposed_num,shift FROM xray_info  WHERE report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(xray_info_sql)
xray_info_ob = cursor.fetchall()
print(xray_info_ob)

for station, product_model, service_time, location, exposed_num, shift in xray_info_ob:
    row_cells = table.add_row().cells
    row_cells[0].text = station
    row_cells[1].text = product_model
    row_cells[2].text = service_time
    row_cells[3].text = location
    row_cells[4].text = str(exposed_num)  # 整型/浮点的写不进去，以后要写个自动判断程序
    row_cells[5].text = shift

p = document.add_paragraph()
run = p.add_run('{}各站点的X射线行李包检查系统每天24小时工作。有行李在传送带上通过时电极管发射X射线，'.format(unit_name))
run.font.size = Pt(11)
power_light_sql = "SELECT station,location FROM xray_info WHERE power_light=0 AND report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(power_light_sql)
power_light_ob = cursor.fetchall()
power_light = merger.merger(power_light_ob)
if power_light == '':
    run = p.add_run('各站点行包检测仪的通电指示灯均正常工作，')
else:
    run = p.add_run('除{}行包检测仪的通电指示灯损坏外,其余通电指示灯均正常工作，'.format(power_light))

xray_light_sql = "SELECT station,location FROM xray_info WHERE xray_light=0 AND report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(xray_light_sql)
xray_light_ob = cursor.fetchall()
xray_light = merger.merger(xray_light_ob)
if xray_light == '':
    run = p.add_run('各站点行包检测仪的X射线发射指示灯均正常工作。')
else:
    run = p.add_run('除{}行包检测仪的X射线发射指示灯损坏外,其余X射线发射指示灯均正常工作。'.format(xray_light))

para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('三、样品采集和现场检测')
run.font.size = Pt(12)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('根据卫生标准的要求和现场调查情况，在直线距离不超过5cm，截面积不小于10cm')
run = p.add_run('2')

run.font.superscript = True  # 文字上标
# run.font.subscript=True   # 文字下标

run = p.add_run('的接受面积上进行平均测量，在行李包入口、行李包出口、旅客通过右侧面、行李包系统顶面按梅花布点法进行了检测，检测报告取最大值；同时在工作人员操作台设检测点。')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('四、检测结果与评价')
run.font.size = Pt(12)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('放射工作人员在较长时间内连续或间断受到超当量剂量（50mSv）的外照射，'
                '易发生急性（慢性）放射性皮肤皮炎和急性放射性皮肤、粘膜溃疡；'
                '超过剂量1.5Sv的外照射易损害造血系统，引起白细胞减少；长期接触过量放射照射还有后遗效应。')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('检测结果表明：所检测X线行李包检查系统表面辐射剂量水平符合《X射线行李包检查系统卫生防护标准》（GBZ127－2002）的要求。（具体检测结果见附表）')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('四、检测结果与评价')
run.font.size = Pt(12)
para_format = p.paragraph_format
para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('1.加强行李包检查系统操作人员安全教育，加大职工操作岗位与行李包检查系统之间的距离，及时维修破损的铅帘和设备指示灯。保护职工、旅客身体健康。')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
# para_format.line_spacing = 1.2

p = document.add_paragraph()
run = p.add_run('2.用人单位应组织接触职业性有害因素的劳动者进行上岗前、在岗期间和离岗时的职业健康检查，发现职业禁忌证时，应及时调离。')
run.font.size = Pt(11)
para_format = p.paragraph_format
para_format.first_line_indent = Pt(22)  # 首行缩进
# para_format.line_spacing = 1.2

document.add_page_break()  # 添加分布符

p = document.add_paragraph()
run = p.add_run('附表：')
run.font.size = Pt(11)

station_tittle_sql = "SELECT DISTINCT station FROM xray_info WHERE report_id=(SELECT MAX(report_id) FROM report_info)"
cursor.execute(station_tittle_sql)
station_tittle_ob = cursor.fetchall()
i = 1
for station_tittle in station_tittle_ob:
    print(station_tittle[0])
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('\n表{} {}X射线行李包检查系统检测结果'.format(i, station_tittle[0]))
    run.font.size = Pt(12)
    i += 1
    xray_result_sql = "SELECT location,entrance,exit_door,left_wall,right_wall,top,workbench FROM xray_info WHERE station= %s AND report_id=(SELECT MAX(report_id) FROM report_info)"
    cursor.execute(xray_result_sql, station_tittle[0])
    xray_result_ob = cursor.fetchall()
    col_num = len(xray_result_ob)
    print('这是', xray_result_ob, len(xray_result_ob))
    xray_result=tupleliststr.any_to_str(xray_result_ob)

    headname = ['检测位置', '限值标准（µGy/h）', '检测空气比释动能率最大值 （µGy/h）']
    table = document.add_table(rows=8, cols=2 + col_num, style='Table Grid')
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(0, 1 + col_num))
    headrow_cells = table.rows[0].cells

    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for z in range(3):
        headrow_cells[z].text = headname[z]
    table.cell(2, 0).text = '行李包入口'
    table.cell(3, 0).text = '行李包出口'
    table.cell(4, 0).text = '旅客通过右侧面'
    table.cell(5, 0).text = '旅客通过左侧面'
    table.cell(6, 0).text = '行李包系统顶面'
    table.cell(7, 0).text = '操作人员工作台'
    table.cell(2, 1).text = '≤5'
    table.cell(3, 1).text = '≤5'
    table.cell(4, 1).text = '≤5'
    table.cell(5, 1).text = '≤5'
    table.cell(6, 1).text = '≤5'
    table.cell(7, 1).text = '—'

    for y in range(col_num):
        for x in range(1,8):
            table.cell(x, y+2).text = xray_result[y][x-1]
            table.rows[x].height = Cm(0.8)
    # table.cell(1, y).text = '≤5'
    # table.cell(2, 1).text = '≤5'
    # table.cell(3, 1).text = '≤5'
    # table.cell(4, 1).text = '≤5'
    # table.cell(5, 1).text = '≤5'
    # table.cell(6, 1).text = '≤5'
    # table.cell(7, 1).text = '—'
    tr=table.rows[0]._tr
    trPr = tr.get_or_add_trPr()  # 获取或添加表行属性
    trHeading = OxmlElement('w:tblHeader')  # 设置word表格的行在各页顶端以标题形式重复出现
    trPr.append(trHeading)
    tr=table.rows[1]._tr
    trPr = tr.get_or_add_trPr()  # 获取或添加表行属性
    trHeading = OxmlElement('w:tblHeader')  # 设置word表格的行在各页顶端以标题形式重复出现
    trPr.append(trHeading)


# 添加最后封面的分节符
current_section = document.add_section()
current_section.start_type = WD_SECTION.NEW_PAGE

header = document.sections[3].header
header.is_linked_to_previous = False
footer = document.sections[3].footer
footer.is_linked_to_previous = False

p = document.add_paragraph()  # 添加空白段落
p = document.add_paragraph('检 测 报 告 书 说 明')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(21)

p = document.add_paragraph('一、对检测结果如有异议，请于收到报告之日起15天内向本中心提出。')
p.runs[0].font.size = Pt(10.5)
p = document.add_paragraph(
    '二、委托检测，如委托者自带检品送检，本疾控中心只对检品负责，不对检品来源负责；如受委托抽样，本疾控中心对抽样过程及检品负责。报告书及检测结果不得做鉴定、评优、审批及商品宣传用。')
p.runs[0].font.size = Pt(10.5)
p = document.add_paragraph('三、监督检测，系按照有关法规进行的监督性检测。')
p.runs[0].font.size = Pt(10.5)
p = document.add_paragraph('四、鉴定检测，系对新产品、新工艺、新资源的卫生质量检测。')
p.runs[0].font.size = Pt(10.5)
p = document.add_paragraph('五、本报告书改动无效。')
p.runs[0].font.size = Pt(10.5)
p = document.add_paragraph('六、未经本中心批准，不得复制检测报告。')
p.runs[0].font.size = Pt(10.5)

for i in range(10):
    document.add_paragraph('')

p = document.add_paragraph('通讯地址：济南市车站街165号\t\t\t\t邮编：250001')
p.runs[0].font.size = Pt(12)
p = document.add_paragraph('电\t话：0531-82425177\t\t\t\t传真：0531-82421385')
p.runs[0].font.size = Pt(12)


id_sql="SELECT MAX(report_id) FROM report_info"
cursor.execute(id_sql)
id_ob=cursor.fetchone()
id=id_ob[0]
doc_path = 'D:\Xray\{} {} {}.docx'.format(report_num,unit_name,id)
document.save(doc_path)
